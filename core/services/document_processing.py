from core.database.models import RAGDocument
from core.schemas.enums import ProcessingStatus
from langchain_community.document_loaders import UnstructuredFileLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.schema import Document
import uuid
from pathlib import Path
import pickle
from sqlalchemy import select
from typing import List
from fastapi import HTTPException
from config.logger import app_logger
from core.services.file_upload import FileUploadService
from core.database.session import AsyncSessionLocal
from core.schemas.enums import FileType
import chardet
import hashlib
from unstructured.partition.auto import partition
from unstructured.partition.pdf import partition_pdf
from pi_heif import register_heif_opener
import os
import asyncio
from datetime import datetime
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import text


# PDF/DOCX Processing
from pypdf import PdfReader
from docx import Document as DocxDocument

# Chunking - DÜZELTİLMİŞ IMPORT
from langchain_experimental.text_splitter import SemanticChunker, BreakpointThresholdType

class DocumentProcessor:
    def __init__(self):
        self.embedding_model = HuggingFaceEmbeddings(
            model_name="sentence-transformers/paraphrase-multilingual-mpnet-base-v2"
        )
        self.processed_dir = Path("processed_documents")
        self.processed_dir.mkdir(exist_ok=True)

    async def process_document(
        self,
        file_path: str,
        text_content: str,
        user_id: str,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        max_file_size: int = 50 * 1024 * 1024
    ) -> List[RAGDocument]:
        """
        Güncellenmiş belge işleme akışı:
        1. Doğrudan gelen içeriği kullan
        2. Semantik chunking
        3. Veritabanına kayıt
        """
        try:
            # 1. Gelen içeriği temizle
            #cleaned_content = self._clean_text(text_content)
            
            # 2. Semantik Chunking
            chunks = await self._semantic_chunking(
                text_content,
                chunk_size,
                chunk_overlap
            )
            
            # 3. Veritabanına Toplu Kayıt
            return await self._bulk_insert_chunks(
                file_path=file_path,
                user_id=user_id,
                chunks=chunks,
                original_content=text_content
            )
            
        except Exception as e:
            app_logger.error(f"Belge işleme hatası: {str(e)}")
            raise HTTPException(500, "Belge işlenemedi")

    def _validate_file(self, file_path: str, max_size: int):
        """Dosya validasyonu"""
        # Tip kontrolü ekleyelim
        if not isinstance(file_path, (str, bytes, os.PathLike)):
            raise TypeError(f"Geçersiz dosya yolu tipi: {type(file_path)}")
        
        if not os.path.exists(file_path):
            raise HTTPException(404, "Dosya bulunamadı")
            
        if os.path.getsize(file_path) > max_size:
            raise HTTPException(413, "Dosya boyutu limiti aşıldı (50MB)")

    async def _extract_text(self, file_path: str) -> str:
        """Dosya türüne göre metin çıkarımı"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == ".pdf":
                text = await self._read_pdf(file_path)
            elif file_ext == ".docx":
                text = await self._read_docx(file_path)
            else:
                raise HTTPException(400, "Desteklenmeyen dosya formatı")
            
            
            # Metni temizleme
            return self._clean_text(text)

        except Exception as e:
            app_logger.error(f"Metin çıkarım hatası: {str(e)}")
            raise HTTPException(500, "Dosya içeriği okunamadı")


    # PDF'den metin çıkarımı

    async def _read_pdf(self, file_path: str) -> str:
        """PDF'den metin çıkarımı (Unstructured ile, OCR devre dışı)"""
        from unstructured.partition.pdf import partition_pdf

        elements = partition_pdf(
            filename=file_path,
            strategy="hi_res",         # "auto" yerine "hi_res" veya "fast" vb. kullanılabilir
            infer_table_structure=True,
            languages=["tur", "eng"],
            ocr_strategy="none"        # <-- OCR kapalı
        )
        return "\n\n".join([str(e) for e in elements])


    async def _read_pdf(self, file_path: str) -> str:
        """PDF'den metin çıkarımı (Unstructured ile)"""
        elements = partition_pdf(
            filename=file_path,
            strategy="auto",
            infer_table_structure=True,
            languages=["tur", "eng"]
        )
        return "\n\n".join([str(e) for e in elements])

    async def _read_docx(self, file_path: str) -> str:
        """DOCX'ten metin çıkarımı"""
        def sync_read():
            doc = DocxDocument(file_path)
            return "\n".join([p.text for p in doc.paragraphs])
            
        return await asyncio.to_thread(sync_read)

    async def _semantic_chunking(
        self,
        text: str,
        chunk_size: int,
        chunk_overlap: int
    ) -> List[str]:
        """Advanced semantic chunking with multiple threshold strategies"""
        chunker = SemanticChunker(
            embeddings=self.embedding_model,
            breakpoint_threshold_type=BreakpointThresholdType.PERCENTILE,  # Default strategy
            breakpoint_threshold_amount=85,  # 85th percentile cutoff
            min_chunk_size=100  # Minimum characters per chunk
        )
        
        return await asyncio.to_thread(
            chunker.split_text,
            text=text,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )

    # Optional alternative strategies
    def _create_standard_deviation_chunker(self):
        return SemanticChunker(
            embeddings=self.embedding_model,
            breakpoint_threshold_type=BreakpointThresholdType.STANDARD_DEVIATION,
            breakpoint_threshold_amount=2.5  # 2.5 standard deviations
        )

    def _create_interquartile_chunker(self):
        return SemanticChunker(
            embeddings=self.embedding_model,
            breakpoint_threshold_type=BreakpointThresholdType.INTERQUARTILE,
            breakpoint_threshold_amount=1.8  # 1.8 * IQR
        )

    def _create_gradient_chunker(self):
        return SemanticChunker(
            embeddings=self.embedding_model,
            breakpoint_threshold_type="gradient",  # Experimental gradient method
            breakpoint_threshold_amount=90  # 90th percentile on gradient
        )

    async def _bulk_insert_chunks(
        self,
        file_path: str,
        user_id: str,
        chunks: List[str],
        original_content: str
    ) -> List[RAGDocument]:
        """Toplu veri ekleme ile performans optimizasyonu"""
        file_checksum = hashlib.sha256(original_content.encode()).hexdigest()
        file_size = os.path.getsize(file_path)
        file_type = os.path.splitext(file_path)[1][1:].upper()
        created_at = datetime.utcnow()

        # RAGDocument nesnelerini oluştur
        documents = [
            RAGDocument(
                title=f"{os.path.basename(file_path)} - Chunk {i}",
                content=self._clean_text(chunk),
                user_id=user_id,
                file_path=file_path,
                file_type=file_type,
                file_size=file_size,
                file_checksum=file_checksum,
                processing_status=ProcessingStatus.PROCESSED,
                chunk_size=len(chunk),
                chunk_index=i,
                created_at=created_at,
                updated_at=created_at,
                metadata={
                    "chunking_method": "semantic_v2",
                    "language": "turkish",
                    "checksum": hashlib.sha256(chunk.encode()).hexdigest()
                }
            )
            for i, chunk in enumerate(chunks)
        ]

        # Toplu veritabanı işlemi
        async with AsyncSessionLocal() as session:
            try:
                session.add_all(documents)
                await session.commit()
                return documents
            except Exception as e:
                await session.rollback()
                app_logger.error(f"Veritabanı kayıt hatası: {str(e)}")
                raise HTTPException(500, "Veritabanı işlemi başarısız")

    def _detect_file_type(self, filename: str) -> str:
        ext = filename.split(".")[-1].lower()
        if ext in ["pdf", "docx", "txt", "md"]:
            return ext
        raise ValueError("Desteklenmeyen dosya formatı")

    def _move_to_processed(self, temp_path: str, doc_id: str) -> Path:
        final_path = self.processed_dir / f"{doc_id}.pdf"
        Path(temp_path).rename(final_path)
        return final_path 

    async def get_user_documents(self, user_id: str):
        result = await self.db.execute(
            select(RAGDocument)
            .filter(RAGDocument.user_id == user_id)
            .order_by(RAGDocument.created_at.desc())
        )
        return result.scalars().all()

    async def _store_embeddings(self, doc: RAGDocument, chunks: List[Document]):
        try:
            from core.rag.vector_store import get_vector_store
            vector_store = get_vector_store()
            
            texts = [chunk.page_content for chunk in chunks]
            embeddings = self.embedding_model.embed_documents(texts)
            
            await vector_store.bulk_insert(
                texts=texts,
                embeddings=embeddings,
                metadatas=[{
                    "doc_id": doc.id,
                    "source": doc.title,
                    "page": idx
                } for idx in range(len(chunks))]
            )
            
            doc.processing_status = "processed"
            await self.db.commit()
        except Exception as e:
            doc.processing_status = "failed"
            doc.processing_errors = str(e)
            await self.db.commit()
            raise 

    async def _save_embeddings(self, chunks: List[Document]):
        try:
            # Numpy array'i bytes'a çevirme
            embeddings = self.embedding_model.embed_documents([chunk.page_content for chunk in chunks])
            return [emb.tobytes() for emb in embeddings]  # Numpy array -> bytes
        except Exception as e:
            app_logger.error(f"Embedding hatası: {str(e)}")
            raise 

    async def read_file_content(self, content: bytes, file_extension: str) -> str:
        """Dosya içeriğini doğru encoding ile okur"""
        try:
            # 1. Encoding tespiti
            detected = chardet.detect(content)
            encoding = detected['encoding'] or 'utf-8'
            
            # 2. Türkçe karakterler için özel handling
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                return content.decode('iso-8859-1')  # Fallback encoding
            except LookupError:
                return content.decode('utf-8', errors='replace')
            
        except Exception as e:
            app_logger.warning(f"Dosya içeriği okunamadı ({file_extension}): {str(e)}")
            return "" 


    def _clean_text(self, text: str) -> str:
        """Metni temizleme ve normalleştirme"""
        return text.replace('\x00', '').replace('\ufeff', '').strip()

    async def _adaptive_chunking(self, text: str) -> List[str]:
        """Bağlama duyarlı chunklama"""
       # from semantic_chunking import SemanticChunker
        chunker = SemanticChunker(
            breakpoint_threshold_type="percentile",
            breakpoint_threshold_amount=85
        )
        return chunker.split_text(text)

    async def _batch_embedding(self, chunks: List[str]) -> List[List[float]]:
        """Toplu embedding işlemi"""
        from sentence_transformers import SentenceTransformer
        model = SentenceTransformer('sentence-transformers/paraphrase-multilingual-mpnet-base-v2')
        return model.encode(chunks).tolist()

    def _generate_semantic_hash(self, text: str) -> str:
        """Metin için anlamsal hash üretimi"""
        return hashlib.sha256(text.encode('utf-8')).hexdigest()[:16] 